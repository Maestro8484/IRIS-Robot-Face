from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins (compact for print)
for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(9)

def set_font(run, bold=False, italic=False, size=9, color=None):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    if level == 1:
        set_font(run, bold=True, size=13, color=(30, 100, 200))
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1E64C8')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        set_font(run, bold=True, size=10, color=(50, 50, 50))
    elif level == 3:
        set_font(run, bold=True, italic=True, size=9, color=(100, 100, 100))
    return p

def add_body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    set_font(run, size=9)
    return p

def add_tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.4)
    r1 = p.add_run("Tip: ")
    set_font(r1, bold=True, size=8.5, color=(20, 130, 60))
    r2 = p.add_run(text)
    set_font(r2, italic=True, size=8.5, color=(20, 130, 60))

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        shade_cell(hrow.cells[i], 'D0DCFA')
        p = hrow.cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(h)
        set_font(run, bold=True, size=8.5, color=(30, 60, 140))

    # data rows
    for ri, row_data in enumerate(rows):
        drow = t.rows[ri + 1]
        fill = 'F5F7FF' if ri % 2 == 1 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            shade_cell(drow.cells[ci], fill)
            p = drow.cells[ci].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(val)
            set_font(run, bold=(ci == 0), size=8.5)

    if col_widths:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = col_widths[i]

    for row in t.rows:
        row.height = Cm(0.52)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(1)
    return t

def add_pipeline():
    lines = [
        ("You say the wake word",                                              False),
        ("  -> brief audio flush (~0.15s)",                                   False),
        ("  -> IRIS records you speaking",                                     False),
        ("  -> you stop talking",                                              False),
        ("  -> IRIS waits  <-- BIGGEST TUNABLE DELAY",                        True),
        ("        Audio tab: Wait after silence  (default 1.5s, try 0.7s)",   True),
        ("  -> audio sent to Whisper for speech-to-text",                     False),
        ("  -> text sent to AI for an answer",                                False),
        ("  -> answer spoken aloud via Chatterbox",                           False),
    ]
    for text, highlight in lines:
        pp = doc.add_paragraph()
        pp.paragraph_format.space_before = Pt(0)
        pp.paragraph_format.space_after  = Pt(0)
        pp.paragraph_format.left_indent  = Cm(0.8)
        run = pp.add_run(text)
        set_font(run, bold=highlight, size=8.5,
                 color=(170, 60, 0) if highlight else (40, 40, 40))

# ── TITLE ─────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(2)
r = p.add_run("IRIS  —  Web UI Settings Guide")
set_font(r, bold=True, size=16, color=(20, 60, 180))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(8)
r2 = p2.add_run("http://192.168.1.200:5000   |   Changes save instantly.  Use Persist to SD to survive reboot.")
set_font(r2, italic=True, size=8.5, color=(120, 120, 120))

# ── AUDIO ──────────────────────────────────────────────────────────────────────
add_heading("Audio tab", 1)
add_body("Controls how IRIS listens after you say the wake word.")
add_heading("Adult", 2)
add_table(
    ["Setting", "What it does", "Default", "Range"],
    [
        ["Max record time",    "How long IRIS waits for you to finish talking before giving up.",                                         "10s",  "1-60s"],
        ["Wait after silence", "How long IRIS waits after you stop talking before sending your words for processing. Lower = faster.",    "1.5s", "0.1-10s"],
        ["Mic sensitivity",    "How quiet the room must be before IRIS considers you done. Raise this in noisy rooms.",                  "300",  "50-5000"],
    ],
    [Cm(3.8), Cm(8.6), Cm(1.6), Cm(1.8)]
)
add_tip("Reducing 'Wait after silence' from 1.5s to 0.7s cuts ~0.8s off every response — the single biggest latency win.")

add_heading("Kids Mode", 2)
add_body("Same three settings with more forgiving defaults — longer waits, quieter mic gate.")
add_table(
    ["Setting", "Default"],
    [
        ["Max record time",    "14s"],
        ["Wait after silence", "3.5s"],
        ["Mic sensitivity",    "150"],
    ],
    [Cm(5), Cm(4)]
)

# ── WAKE WORD ──────────────────────────────────────────────────────────────────
add_heading("Wake Word tab", 1)
add_body("Controls how IRIS detects the wake word.")
add_table(
    ["Setting", "What it does", "Default", "Range"],
    [
        ["Detection confidence", "How sure IRIS needs to be before waking up. Higher = fewer false triggers, may miss quiet wake words.", "0.90",  "0.5-1.0"],
        ["Startup delay",        "Brief audio flush after wake word fires so it does not bleed into your question. Rarely needs changing.", "0.15s", "0.05-1.0s"],
    ],
    [Cm(3.8), Cm(8.6), Cm(1.6), Cm(1.8)]
)

# ── VOICE ──────────────────────────────────────────────────────────────────────
add_heading("Voice tab", 1)
add_body("Controls how IRIS speaks. Chatterbox (on GandalfAI) is the primary voice; Piper is the offline fallback.")
add_table(
    ["Setting", "What it does", "Default"],
    [
        ["Chatterbox on/off", "Turn the main voice engine on or off. Off = falls back to simpler Piper voice.", "On"],
        ["Reference voice",   "Audio file used to clone IRIS's voice. Swap this to change how IRIS sounds.",    "iris_voice.wav"],
        ["Expressiveness",    "0 = flat/robotic    0.45 = dry wit (default)    1.0 = theatrical",              "0.45"],
    ],
    [Cm(3.8), Cm(9.4), Cm(2.8)]
)

# ── CONVERSATION ───────────────────────────────────────────────────────────────
add_heading("Conversation tab", 1)
add_body("Controls back-and-forth conversation behaviour.")
add_heading("Follow-up & Context", 2)
add_table(
    ["Setting", "What it does", "Default", "Range"],
    [
        ["Follow-up window",      "After answering, how long IRIS keeps listening for a follow-up before going back to sleep.",  "2s",    "1-60s"],
        ["Kids follow-up window", "Same, for kids mode.",                                                                        "15s",   "1-120s"],
        ["Max follow-up turns",   "Back-and-forth exchanges before IRIS forgets the conversation and starts fresh.",             "3",     "1-20"],
        ["Memory timeout",        "If nobody talks to IRIS for this long, it forgets the conversation context.",                "5 min", "30s-1hr"],
    ],
    [Cm(3.8), Cm(8.4), Cm(1.6), Cm(2.2)]
)
add_heading("Response Length", 2)
add_body("IRIS auto-picks how long to make its answer based on the question. No restart needed.")
add_table(
    ["Tier", "When IRIS uses it", "Default length"],
    [
        ["Short",        "Yes/no, greetings, quick facts",             "~2-3 sentences"],
        ["Medium",       "Explanations, how-to questions",             "~4-6 sentences"],
        ["Long",         "Stories, step-by-step guides, lists",        "~8-12 sentences"],
        ["Max",          "Tell me everything about..., essays",        "~15+ sentences"],
        ["Voice cutoff", "Maximum spoken before Chatterbox cuts off",  "~5-8 sentences"],
    ],
    [Cm(2.6), Cm(8.6), Cm(4.8)]
)

# ── SLEEP ──────────────────────────────────────────────────────────────────────
add_heading("Sleep tab", 1)
add_body("Controls sleep/wake behaviour and display brightness at night.  Auto-schedule: Sleep 9 PM / Wake 7:30 AM.")
add_heading("Mouth Display Brightness", 2)
add_table(
    ["Setting", "What it does", "Default", "Range"],
    [
        ["Awake brightness", "Mouth display brightness during the day.",               "8", "0-15"],
        ["Sleep brightness", "Mouth display brightness at night. Keep this very dim.", "1", "0-15"],
    ],
    [Cm(3.8), Cm(8.4), Cm(1.6), Cm(2.2)]
)
add_body("'Save & Apply Now' changes brightness immediately without a restart.")
add_heading("Sleep LED Glow  (requires assistant restart)", 2)
add_table(
    ["Setting", "Default"],
    [
        ["Peak brightness", "26  (very dim by design)"],
        ["Floor brightness", "3"],
        ["Pulse speed",      "8s per cycle"],
    ],
    [Cm(4), Cm(12)]
)

# ── LIGHTS ─────────────────────────────────────────────────────────────────────
add_heading("Lights tab  (requires assistant restart)", 1)
add_body("LED glow colour and pulse while IRIS is awake.")
add_table(
    ["Setting", "Normal — cyan", "Kids mode — yellow"],
    [
        ["Peak brightness", "65",  "62"],
        ["Floor brightness", "3",  "—"],
        ["Pulse speed",      "5s", "4s"],
    ],
    [Cm(4), Cm(5), Cm(5)]
)

# ── GANDALF AI ─────────────────────────────────────────────────────────────────
add_heading("Gandalf AI tab", 1)
add_body("Controls which AI brain IRIS uses to answer questions.")
add_table(
    ["Setting", "What it does", "Default"],
    [
        ["Adult model", "AI model used for adult conversations.", "iris"],
        ["Kids model",  "AI model used when kids mode is active.", "iris-kids"],
    ],
    [Cm(3.2), Cm(9.4), Cm(3.4)]
)

# ── SYSTEM ─────────────────────────────────────────────────────────────────────
add_heading("System tab", 1)
add_heading("Speaker Volume", 2)
add_table(
    ["Setting", "What it does", "Default", "Range"],
    [
        ["Volume slider",  "IRIS's speaker volume right now. Takes effect immediately.",                "121", "0-127"],
        ["Volume ceiling", "Maximum volume that voice commands (turn it up) are allowed to reach.",    "127", "60-127"],
    ],
    [Cm(3.2), Cm(9), Cm(1.6), Cm(2.2)]
)
add_heading("SD Persistence", 2)
add_table(
    ["Status bar colour", "Meaning"],
    [
        ["Green", "Saved to SD card — safe on reboot."],
        ["Amber", "RAM only — will be lost on reboot. Click Persist to SD."],
    ],
    [Cm(4), Cm(12)]
)

# ── PIPELINE ───────────────────────────────────────────────────────────────────
add_heading("How a response happens", 1)
add_body("Every voice interaction follows this chain. The Bench tab shows real timings for each stage.")
add_pipeline()

# ── SAVE ───────────────────────────────────────────────────────────────────────
out = r'C:\Users\SuperMaster\Documents\PlatformIO\IRIS-Robot-Face\GUIDE-settings.docx'
doc.save(out)
print("saved:", out)
